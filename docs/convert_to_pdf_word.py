#!/usr/bin/env python3
"""
Automatically convert TAG_DEPLOYMENT_GUIDE.md to PDF and Word documents

This script generates both PDF and Word (.docx) files automatically
without requiring any user interaction.

Usage:
    python3 convert_to_pdf_word.py
"""

import os
import sys
import subprocess
from pathlib import Path

def install_package(package):
    """Install a Python package"""
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package, "--upgrade"],
                             stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except subprocess.CalledProcessError:
        return False

def convert_to_word(md_file, docx_file):
    """Convert markdown to Word document using python-docx"""
    print("Creating Word document...")
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
    except ImportError:
        print("  Installing python-docx...")
        if not install_package("python-docx"):
            print("  ✗ Failed to install python-docx")
            return False
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import re
    
    try:
        # Read markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Create Word document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Parse markdown content
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_block_lines = []
        code_block_lang = ''
        
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Handle code blocks
            if stripped.startswith('```'):
                if in_code_block:
                    # End of code block
                    code_text = '\n'.join(code_block_lines)
                    para = doc.add_paragraph()
                    run = para.add_run(code_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    para.style = 'No Spacing'
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                    in_code_block = False
                    code_block_lines = []
                else:
                    # Start of code block
                    in_code_block = True
                    code_block_lang = stripped[3:] if len(stripped) > 3 else ''
                i += 1
                continue
            
            if in_code_block:
                code_block_lines.append(line)
                i += 1
                continue
            
            # Skip empty lines (but preserve spacing between sections)
            if not stripped:
                if i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].strip().startswith('#'):
                    doc.add_paragraph()
                i += 1
                continue
            
            # Headers
            if stripped.startswith('# '):
                heading = doc.add_heading(stripped[2:], level=1)
                for run in heading.runs:
                    run.font.size = Pt(24)
            elif stripped.startswith('## '):
                heading = doc.add_heading(stripped[3:], level=2)
                for run in heading.runs:
                    run.font.size = Pt(20)
            elif stripped.startswith('### '):
                heading = doc.add_heading(stripped[4:], level=3)
                for run in heading.runs:
                    run.font.size = Pt(16)
            elif stripped.startswith('#### '):
                heading = doc.add_heading(stripped[5:], level=4)
                for run in heading.runs:
                    run.font.size = Pt(14)
            # Horizontal rule
            elif stripped.startswith('---'):
                para = doc.add_paragraph('─' * 50)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Lists
            elif stripped.startswith('- ') or stripped.startswith('* '):
                para = doc.add_paragraph(style='List Bullet')
                add_formatted_text(para, stripped[2:])
            elif re.match(r'^\d+\.\s', stripped):
                para = doc.add_paragraph(style='List Number')
                add_formatted_text(para, re.sub(r'^\d+\.\s', '', stripped))
            # Regular paragraphs
            else:
                para = doc.add_paragraph()
                add_formatted_text(para, stripped)
            
            i += 1
        
        # Save document
        doc.save(docx_file)
        file_size = docx_file.stat().st_size / 1024
        print(f"  ✓ Word document created: {docx_file.name} ({file_size:.1f} KB)")
        return True
        
    except Exception as e:
        print(f"  ✗ Error creating Word document: {e}")
        import traceback
        traceback.print_exc()
        return False

def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, code, italic)"""
    from docx.shared import Pt
    import re
    
    if not text:
        return
    
    # Split by markdown patterns, preserving them
    pattern = r'(\*\*.*?\*\*|`.*?`|\*[^*].*?\*)'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        elif part.startswith('**') and part.endswith('**') and len(part) > 4:
            # Bold
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            # Inline code
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**') and len(part) > 2:
            # Italic (but not bold)
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)

def convert_to_pdf_with_playwright(md_file, pdf_file):
    """Convert markdown to PDF using Playwright (headless browser)"""
    print("Creating PDF using headless browser...")
    
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print("  Installing playwright...")
        if not install_package("playwright"):
            return False
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            return False
        # Install playwright browsers
        print("  Installing playwright browsers (this may take a moment, ~160MB)...")
        try:
            result = subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"],
                                  capture_output=True, text=True)
            if result.returncode != 0:
                print(f"  Warning: Browser installation had issues. Trying to continue...")
        except Exception as e:
            print(f"  Warning: Could not install browsers automatically: {e}")
            print("  You may need to run: python3 -m playwright install chromium")
    
    try:
        import markdown
    except ImportError:
        if not install_package("markdown"):
            return False
        import markdown
    
    try:
        # Read and convert markdown to HTML
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        md = markdown.Markdown(extensions=['extra', 'codehilite', 'tables', 'toc'])
        html_content = md.convert(md_content)
        
        # Create styled HTML
        html_with_style = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>3E Tag Manager Deployment Guide</title>
            <style>
                @page {{
                    size: letter;
                    margin: 0.75in;
                }}
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    color: #333;
                    max-width: 900px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                h1 {{
                    font-size: 24pt;
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                    margin-top: 0;
                    margin-bottom: 20px;
                }}
                h2 {{
                    font-size: 20pt;
                    color: #34495e;
                    margin-top: 25px;
                    margin-bottom: 15px;
                    border-bottom: 2px solid #ecf0f1;
                    padding-bottom: 5px;
                }}
                h3 {{
                    font-size: 16pt;
                    color: #34495e;
                    margin-top: 20px;
                    margin-bottom: 10px;
                }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 6px;
                    border-radius: 3px;
                    font-family: "Courier New", monospace;
                    font-size: 10pt;
                }}
                pre {{
                    background-color: #f8f8f8;
                    border: 1px solid #ddd;
                    border-left: 4px solid #3498db;
                    padding: 12px;
                    overflow-x: auto;
                    font-size: 10pt;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}
                ul, ol {{
                    margin: 10px 0;
                    padding-left: 30px;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 15px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #f2f2f2;
                    font-weight: bold;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Use Playwright to generate PDF
        with sync_playwright() as p:
            try:
                browser = p.chromium.launch()
            except Exception as e:
                # Browser not found, try to install
                print("  Browser not found. Installing...")
                subprocess.run([sys.executable, "-m", "playwright", "install", "chromium"])
                browser = p.chromium.launch()
            
            page = browser.new_page()
            page.set_content(html_with_style)
            page.pdf(path=str(pdf_file), format='Letter', margin={
                'top': '0.75in',
                'right': '0.75in',
                'bottom': '0.75in',
                'left': '0.75in'
            })
            browser.close()
        
        file_size = pdf_file.stat().st_size / 1024
        print(f"  ✓ PDF created: {pdf_file.name} ({file_size:.1f} KB)")
        return True
        
    except Exception as e:
        print(f"  ✗ Error creating PDF: {e}")
        return False

def main():
    script_dir = Path(__file__).parent
    md_file = script_dir / "TAG_DEPLOYMENT_GUIDE.md"
    docx_file = script_dir / "TAG_DEPLOYMENT_GUIDE.docx"
    pdf_file = script_dir / "TAG_DEPLOYMENT_GUIDE.pdf"
    
    if not md_file.exists():
        print(f"Error: {md_file} not found")
        sys.exit(1)
    
    print(f"\nConverting {md_file.name} to Word and PDF...\n")
    
    # Remove old files
    for old_file in [docx_file, pdf_file]:
        if old_file.exists():
            print(f"Removing old {old_file.name}...")
            old_file.unlink()
    
    # Convert to Word
    word_success = convert_to_word(md_file, docx_file)
    
    # Convert to PDF
    pdf_success = convert_to_pdf_with_playwright(md_file, pdf_file)
    
    print("\n" + "="*60)
    if word_success:
        print(f"✓ Word document: {docx_file.name}")
    else:
        print("✗ Word document: Failed")
    
    if pdf_success:
        print(f"✓ PDF document: {pdf_file.name}")
    else:
        print("✗ PDF document: Failed")
    
    if word_success or pdf_success:
        print("\n✓ Conversion complete!")
    else:
        print("\n✗ Conversion failed. Check errors above.")
        sys.exit(1)

if __name__ == "__main__":
    main()
